import logging
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
from docx import Document
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Настройка логирования
logger = logging.getLogger()
logger.setLevel(logging.DEBUG)

# Путь для сохранения сертификатов
BASE_DIR = Path(__file__).resolve().parent.parent
CERTIFICATE_DIR = BASE_DIR / "static" / "certificates"  # Основная папка для сертификатов

# Подкаталоги для различных форматов
IMAGES_DIR = CERTIFICATE_DIR / "images"
PDF_DIR = CERTIFICATE_DIR / "pdf"
WORD_DIR = CERTIFICATE_DIR / "word"

# Путь к файлам изображений и шрифтам
IMAGE_TEMPLATE_PATH = IMAGES_DIR / "certificate.png"  # Шаблон сертификата
SIGNATURE_DIRECTOR_PATH = IMAGES_DIR / "podpis_1.png"  # Подпись директора
SIGNATURE_METHODIST_PATH = IMAGES_DIR / "podpis_2.png"  # Подпись методиста
FONT_PATH = Path(__file__).resolve().parent.parent / "static" / "fonts" / "OpenSans.ttf"

# FONT_PATH = BASE_DIR / "arial.ttf"  # Или шрифт по умолчанию, если arial.ttf не найден


def load_font(font_path, size):
    try:
        logger.debug(f"Загружаем шрифт: {font_path} размер {size}")
        return ImageFont.truetype(str(font_path), size)
    except IOError:
        logger.warning(f"Шрифт {font_path} не найден, используется стандартный.")
        return ImageFont.load_default()


def check_and_create_directories():
    """Проверяет и создает все необходимые папки для сертификатов."""
    try:
        CERTIFICATE_DIR.mkdir(parents=True, exist_ok=True)
        IMAGES_DIR.mkdir(parents=True, exist_ok=True)
        PDF_DIR.mkdir(parents=True, exist_ok=True)
        WORD_DIR.mkdir(parents=True, exist_ok=True)
        logger.debug("Директории для сертификатов проверены и созданы при необходимости.")
    except Exception as e:
        logger.error(f"Ошибка при создании директорий: {e}")
        raise


def save_certificate_image_in_formats(image, base_path):
    try:
        # Сохраняем в форматах .png, .jpg, .bmp
        png_path = base_path / "certificate_completed.png"
        jpg_path = base_path / "certificate_completed.jpg"
        bmp_path = base_path / "certificate_completed.bmp"
        
        # Сохраняем изображение в трех форматах
        image.save(png_path, format="PNG")
        image.save(jpg_path, format="JPEG")
        image.save(bmp_path, format="BMP")

        logger.info(f"Сертификат сохранен в формате PNG: {png_path}")
        logger.info(f"Сертификат сохранен в формате JPG: {jpg_path}")
        logger.info(f"Сертификат сохранен в формате BMP: {bmp_path}")

        return png_path, jpg_path, bmp_path

    except Exception as e:
        logger.error(f"Ошибка при сохранении изображения в разных форматах: {e}")
        return None, None, None


# Генерация сертификата в формате PNG
def generate_png_certificate(student_name, event_name, event_date, issuer_name,
                              director_name, director_position, methodist_name, methodist_position):
    logger.debug("Начинаем генерацию сертификата в формате PNG...")
    try:
        check_and_create_directories()  # Проверка и создание директорий перед генерацией

        image = Image.open(IMAGE_TEMPLATE_PATH)
        signature_director = Image.open(SIGNATURE_DIRECTOR_PATH).resize((200, 100))
        signature_methodist = Image.open(SIGNATURE_METHODIST_PATH).resize((200, 100))

        font_large = load_font(FONT_PATH, 52)
        font_medium = load_font(FONT_PATH, 38)
        font_small = load_font(FONT_PATH, 32)

        draw = ImageDraw.Draw(image)
        # Добавление текста на сертификат
        logger.debug(f"Добавление текста: {student_name}, {event_name}, {event_date}, {issuer_name}")
        draw.text((700, 600), student_name, fill="black", font=font_large)
        draw.text((650, 700), event_name, fill="black", font=font_medium)
        draw.text((480, 1230), event_date, fill="black", font=font_medium)
        draw.text((1060, 1230), issuer_name, fill="black", font=font_medium)
        draw.text((600, 1100), f"{director_name}\n{director_position}", fill="black", font=font_small)
        draw.text((1150, 1100), f"{methodist_name}\n{methodist_position}", fill="black", font=font_small)

        # Наложение подписей
        image.paste(signature_director, (650, 930), signature_director)
        image.paste(signature_methodist, (1200, 930), signature_methodist)

        # Сохраняем сертификат в разных форматах
        save_certificate_image_in_formats(image, IMAGES_DIR)

        output_path = IMAGES_DIR / "certificate_completed.png"
        image.save(output_path)

        logger.info(f"Сертификат PNG сохранен по пути {output_path}")
        return str(output_path)

    except Exception as e:
        logger.error(f"Ошибка при генерации сертификата PNG: {e}")
        return None


def generate_jpg_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position):
    # Генерация изображения
    png_path = generate_png_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
    jpg_path = png_path.replace(".png", ".jpg")

    try:
        with Image.open(png_path) as img:
            rgb_img = img.convert("RGB")
            rgb_img.save(jpg_path, "JPEG")
        return jpg_path
    except Exception as e:
        logging.error(f"Ошибка при преобразовании в JPG: {e}")
        return None


def generate_bmp_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position):
    # Генерация изображения
    png_path = generate_png_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
    bmp_path = png_path.replace(".png", ".bmp")

    try:
        with Image.open(png_path) as img:
            bmp_img = img.convert("RGB")  # BMP также не поддерживает RGBA
            bmp_img.save(bmp_path, "BMP")
        return bmp_path
    except Exception as e:
        logging.error(f"Ошибка при преобразовании в BMP: {e}")
        return None


# Генерация сертификата в формате PDF
def generate_pdf_certificate(student_name, event_name, event_date, issuer_name,
                              director_name, director_position, methodist_name, methodist_position):
    logger.debug("Начинаем генерацию сертификата в формате PDF...")
    try:
        check_and_create_directories()  # Проверка и создание директорий перед генерацией

        # Путь к изображению готового сертификата
        certificate_image_path = IMAGES_DIR / "certificate_completed.png"
        
        if not certificate_image_path.exists():
            logger.error(f"Не найдено изображение сертификата: {certificate_image_path}")
            raise FileNotFoundError(f"Не найдено изображение сертификата: {certificate_image_path}")

        # Путь для сохранения PDF
        output_path = PDF_DIR / "certificate_completed.pdf"

        c = canvas.Canvas(str(output_path), pagesize=letter)

        # Добавляем изображение сертификата
        c.drawImage(str(certificate_image_path), 50, 300, width=500, height=300)

        # Устанавливаем шрифт и размер
        c.setFont("Helvetica", 12)

        # Добавляем текст (дату и подпись руководителя) внизу
        c.drawString(50, 250, f"Дата: {event_date}")
        c.drawString(50, 230, f"Подписано: {director_name}, {director_position}")
        c.drawString(50, 210, f"Методист: {methodist_name}, {methodist_position}")
        c.drawString(50, 190, f"Организатор: {issuer_name}")
        c.drawString(50, 170, f"Участник: {student_name}")
        c.drawString(50, 150, f"Мероприятие: {event_name}")

        c.showPage()
        c.save()

        logger.info(f"Сертификат PDF сохранен по пути {output_path}")
        return str(output_path)

    except Exception as e:
        logger.error(f"Ошибка при генерации PDF сертификата: {e}")
        return None


# Генерация сертификата в формате DOCX
def generate_word_certificate(student_name, event_name, event_date, issuer_name,
                               director_name, director_position, methodist_name, methodist_position):
    logger.debug("Начинаем генерацию сертификата в формате DOCX...")
    try:
        check_and_create_directories()  # Проверка и создание директорий перед генерацией

        # Путь к изображению готового сертификата
        certificate_image_path = IMAGES_DIR / "certificate_completed.png"
        
        if not certificate_image_path.exists():
            logger.error(f"Не найдено изображение сертификата: {certificate_image_path}")
            raise FileNotFoundError(f"Не найдено изображение сертификата: {certificate_image_path}")

        output_path = WORD_DIR / "certificate_completed.docx"
        doc = Document()
        doc.add_heading("Сертификат участника", 0)
        doc.add_paragraph(f"Настоящим подтверждается, что {student_name} принял(а) участие в мероприятии '{event_name}'.")
        doc.add_paragraph(f"Дата проведения: {event_date}")
        doc.add_paragraph(f"Организатор: {issuer_name}")
        doc.add_paragraph(f"Подписано: {director_name}, {director_position}")
        doc.add_paragraph(f"Методист: {methodist_name}, {methodist_position}")

        # Добавляем изображение сертификата
        doc.add_picture(str(certificate_image_path), width=400, height=200)

        doc.save(str(output_path))

        logger.info(f"Сертификат DOCX сохранен по пути {output_path}")
        return str(output_path)

    except Exception as e:
        logger.error(f"Ошибка при генерации Word сертификата: {e}")
        return None


# Главная функция для генерации сертификата в любом из форматов
def generate_certificate_full(student_name, event_name, event_date, issuer_name, director_name,
                              director_position, methodist_name, methodist_position, format_type="png"):
    logger.debug("Начинаем генерацию сертификата...")
    try:
        # Проверка существования файлов шаблонов
        if not IMAGE_TEMPLATE_PATH.exists():
            logger.error(f"Не найден шаблон сертификата: {IMAGE_TEMPLATE_PATH}")
            raise FileNotFoundError(f"Не найден шаблон сертификата: {IMAGE_TEMPLATE_PATH}")

        if not SIGNATURE_DIRECTOR_PATH.exists():
            logger.error(f"Не найдена подпись директора: {SIGNATURE_DIRECTOR_PATH}")
            raise FileNotFoundError(f"Не найдена подпись директора: {SIGNATURE_DIRECTOR_PATH}")
        
        if not SIGNATURE_METHODIST_PATH.exists():
            logger.error(f"Не найдена подпись методиста: {SIGNATURE_METHODIST_PATH}")
            raise FileNotFoundError(f"Не найдена подпись методиста: {SIGNATURE_METHODIST_PATH}")
        
        logger.debug(f"Все необходимые файлы найдены. Формат сертификата: {format_type}")

        # Генерация сертификата в указанном формате
        if format_type == "png":
            return generate_png_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
        elif format_type == "jpg":
            return generate_jpg_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
        elif format_type == "bmp":
            return generate_bmp_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
        elif format_type == "pdf":
            return generate_pdf_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
        elif format_type == "docx":
            return generate_word_certificate(student_name, event_name, event_date, issuer_name, director_name, director_position, methodist_name, methodist_position)
        else:
            raise ValueError("Неверный формат сертификата. Поддерживаемые форматы: png, jpg, bmp, pdf, docx.")
    
    except Exception as e:
        logger.error(f"Ошибка при генерации сертификата: {e}")
        return None
